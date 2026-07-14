"""Shared, reference-matched resume template for the portfolio downloads.

The output intentionally follows a restrained, ATS-friendly one-page format:
centered identity header, compact section rules, right-aligned dates, and plain
black text without decorative tables or color fills.
"""

from __future__ import annotations

import shutil
from pathlib import Path
from typing import Literal, TypedDict


ROOT = Path(__file__).resolve().parents[1]
Locale = Literal["en", "es"]


class Entry(TypedDict):
    organization: str
    location: str
    role: str
    dates: str
    detail: str
    bullets: list[str]


class ResumeContent(TypedDict):
    title: str
    contact: str
    sections: dict[str, str]
    education: list[Entry]
    work: list[Entry]
    projects: list[Entry]
    leadership: list[Entry]
    skills: list[tuple[str, str]]
    training: list[str]


NAME = "JOSE PABLO SAMANO SUAREZ"

CONTENT: dict[Locale, ResumeContent] = {
    "en": {
        "title": "Computer Science Student | Software Engineering | Cybersecurity | Full-Stack Product Systems",
        "contact": (
            "Wolfville, NS / Mexico City | jpss2004@icloud.com | 6195978559 | "
            "linkedin.com/in/jose-pablo-samano-suarez | github.com/jpss2004-bot"
        ),
        "sections": {
            "education": "EDUCATION",
            "work": "WORK EXPERIENCE",
            "projects": "PROJECTS",
            "leadership": "LEADERSHIP EXPERIENCE",
            "skills": "TECHNICAL SKILLS",
            "training": "ADDITIONAL TRAINING",
        },
        "education": [
            {
                "organization": "Acadia University",
                "location": "Wolfville, NS",
                "role": "Bachelor of Science in Computer Science",
                "dates": "Expected May 2027",
                "detail": "",
                "bullets": [
                    "Coursework and project focus: software engineering, data structures, databases, cybersecurity, systems analysis, human-machine interaction, cloud/security fundamentals, and secure application design."
                ],
            },
            {
                "organization": "Trinity College School",
                "location": "Port Hope, ON",
                "role": "High School Diploma",
                "dates": "June 2023",
                "detail": "",
                "bullets": [],
            },
        ],
        "work": [
            {
                "organization": "Cybolt",
                "location": "Mexico City",
                "role": "Cybolt Academy Participant",
                "dates": "May-Sep 2025",
                "detail": "",
                "bullets": [
                    "Completed hands-on workshops on MITRE ATT&CK, ICS/OT security, cyber kill chain, incident response, and defensive blue-team operations.",
                    "Analyzed adversary TTPs and ransomware behavior to connect attack patterns with detection, containment, investigation, and decision-making.",
                    "Built STAR-based incident-response interview cases, translating technical scenarios into clear, evidence-oriented narratives.",
                ],
            },
            {
                "organization": "Michelin",
                "location": "Waterville, NS",
                "role": "Production Operator",
                "dates": "May-Sep 2024",
                "detail": "",
                "bullets": [
                    "Operated and monitored rubber-processing equipment in a high-pressure industrial environment while maintaining quality, safety, and consistency standards.",
                    "Identified and communicated production delays or equipment issues to reduce workflow interruptions and support plant continuity.",
                ],
            },
            {
                "organization": "Legal Shelf",
                "location": "Mexico City",
                "role": "Data Processor",
                "dates": "May-Aug 2023",
                "detail": "",
                "bullets": [
                    "Digitized, archived, and organized legal documents with structured metadata, improving retrieval, traceability, and file administration.",
                    "Handled sensitive information under strict confidentiality and data-protection practices during legal document processing.",
                ],
            },
        ],
        "projects": [
            {
                "organization": "Family Phrase Game",
                "location": "",
                "role": "Developer",
                "dates": "2026",
                "detail": "Flask | HTML | CSS | JavaScript | Render",
                "bullets": [
                    "Built and deployed a Flask web app that turns family-submitted phrases into a playable party game with phrase loading, scoring, and a simple live interface."
                ],
            },
            {
                "organization": "ER Triage & Queue Manager",
                "location": "",
                "role": "System Designer",
                "dates": "2025",
                "detail": "Python | NiceGUI | SQLite",
                "bullets": [
                    "Modeled an emergency-room workflow around intake, vitals, priority scoring, patient state, status history, and dashboard visibility."
                ],
            },
        ],
        "leadership": [
            {
                "organization": "SAVR - Context-Aware Dining Recommendation Platform",
                "location": "Wolfville, NS",
                "role": "Founder / Developer",
                "dates": "2025-Present",
                "detail": "FastAPI | React | TypeScript | SQL",
                "bullets": [
                    "Building a full-stack platform that ranks restaurants using preferences, budget, social context, and user intent.",
                    "Designed Describe Your Night, Build Your Night, and Surprise Me flows that turn structured inputs into prioritized, explainable recommendations.",
                    "Implemented backend and frontend improvements across authentication, restaurant data, recommendation APIs, onboarding, presets, and result components.",
                ],
            }
        ],
        "skills": [
            ("Software", "Python, Java, JavaScript/TypeScript, C, C#, SQL, React, Next.js, FastAPI, Flask, REST APIs, Git"),
            ("Data & Delivery", "PostgreSQL, SQLite, SQLAlchemy, data modeling, validation, testing, Vercel, Render, documentation"),
            ("Cybersecurity", "Incident response, threat analysis, MITRE ATT&CK, cyber kill chain, SOC fundamentals, ICS/OT security, authentication, access control"),
            ("Languages", "Spanish native, English C2, French A2"),
        ],
        "training": [
            "CompTIA Security+ coursework, CompTIA CySA+ coursework, COMP 601 Cybersecurity Fundamentals, and MITRE ATT&CK & Incident Response Workshop.",
            "Ongoing practice in defensive security, incident-response interviews, technical documentation, adversary-behavior analysis, and SOC fundamentals.",
        ],
    },
    "es": {
        "title": "Estudiante de Ciencias de la Computación | Ingeniería de Software | Ciberseguridad | Sistemas Full-Stack",
        "contact": (
            "Wolfville, NS / Ciudad de México | jpss2004@icloud.com | 6195978559 | "
            "linkedin.com/in/jose-pablo-samano-suarez | github.com/jpss2004-bot"
        ),
        "sections": {
            "education": "EDUCACIÓN",
            "work": "EXPERIENCIA LABORAL",
            "projects": "PROYECTOS",
            "leadership": "EXPERIENCIA DE LIDERAZGO",
            "skills": "HABILIDADES TÉCNICAS",
            "training": "FORMACIÓN ADICIONAL",
        },
        "education": [
            {
                "organization": "Acadia University",
                "location": "Wolfville, NS",
                "role": "Bachelor of Science in Computer Science",
                "dates": "Esperado Mayo 2027",
                "detail": "",
                "bullets": [
                    "Enfoque de cursos y proyectos: ingeniería de software, estructuras de datos, bases de datos, ciberseguridad, análisis de sistemas, interacción humano-máquina, fundamentos de nube/seguridad y diseño de aplicaciones seguras."
                ],
            },
            {
                "organization": "Trinity College School",
                "location": "Port Hope, ON",
                "role": "High School Diploma",
                "dates": "Junio 2023",
                "detail": "",
                "bullets": [],
            },
        ],
        "work": [
            {
                "organization": "Cybolt",
                "location": "Ciudad de México",
                "role": "Participante de Cybolt Academy",
                "dates": "May-Sep 2025",
                "detail": "",
                "bullets": [
                    "Completó talleres prácticos sobre MITRE ATT&CK, seguridad ICS/OT, cyber kill chain, respuesta a incidentes y operaciones defensivas de blue team.",
                    "Analizó TTPs adversarias y comportamiento de ransomware para conectar patrones de ataque con detección, contención, investigación y toma de decisiones.",
                    "Construyó casos de entrevista de respuesta a incidentes basados en STAR, traduciendo escenarios técnicos en narrativas claras y orientadas a la evidencia.",
                ],
            },
            {
                "organization": "Michelin",
                "location": "Waterville, NS",
                "role": "Operador de Producción",
                "dates": "May-Sep 2024",
                "detail": "",
                "bullets": [
                    "Operó y monitoreó equipo de procesamiento de hule en un entorno industrial de alta presión, manteniendo estándares de calidad, seguridad y consistencia.",
                    "Identificó y comunicó retrasos de producción o fallas de equipo para reducir interrupciones del flujo de trabajo y apoyar la continuidad de la planta.",
                ],
            },
            {
                "organization": "Legal Shelf",
                "location": "Ciudad de México",
                "role": "Procesador de Datos",
                "dates": "May-Ago 2023",
                "detail": "",
                "bullets": [
                    "Digitalizó, archivó y organizó documentos legales con metadatos estructurados, mejorando la recuperación, la trazabilidad y la administración de archivos.",
                    "Manejó información sensible bajo estrictas prácticas de confidencialidad y protección de datos durante el procesamiento de documentos legales.",
                ],
            },
        ],
        "projects": [
            {
                "organization": "Family Phrase Game",
                "location": "",
                "role": "Desarrollador",
                "dates": "2026",
                "detail": "Flask | HTML | CSS | JavaScript | Render",
                "bullets": [
                    "Construyó y desplegó una aplicación web en Flask que convierte frases familiares en un juego con carga de frases, puntuación y una interfaz en vivo sencilla."
                ],
            },
            {
                "organization": "ER Triage & Queue Manager",
                "location": "",
                "role": "Diseñador de Sistema",
                "dates": "2025",
                "detail": "Python | NiceGUI | SQLite",
                "bullets": [
                    "Modeló un flujo de sala de emergencias en torno a admisión, signos vitales, puntuación de prioridad, estado del paciente, historial de estados y visibilidad en dashboard."
                ],
            },
        ],
        "leadership": [
            {
                "organization": "SAVR - Plataforma de Recomendación Gastronómica Contextual",
                "location": "Wolfville, NS",
                "role": "Fundador / Desarrollador",
                "dates": "2025-Presente",
                "detail": "FastAPI | React | TypeScript | SQL",
                "bullets": [
                    "Construye una plataforma full-stack que clasifica restaurantes según preferencias, presupuesto, contexto social e intención del usuario.",
                    "Diseñó los flujos Describe Your Night, Build Your Night y Surprise Me para convertir entradas estructuradas en recomendaciones priorizadas y explicables.",
                    "Implementó mejoras de backend y frontend en autenticación, datos de restaurantes, APIs de recomendación, onboarding, presets y componentes de resultados.",
                ],
            }
        ],
        "skills": [
            ("Software", "Python, Java, JavaScript/TypeScript, C, C#, SQL, React, Next.js, FastAPI, Flask, REST APIs, Git"),
            ("Datos y entrega", "PostgreSQL, SQLite, SQLAlchemy, modelado de datos, validación, pruebas, Vercel, Render, documentación"),
            ("Ciberseguridad", "Respuesta a incidentes, análisis de amenazas, MITRE ATT&CK, cyber kill chain, fundamentos de SOC, seguridad ICS/OT, autenticación, control de acceso"),
            ("Idiomas", "Español nativo, Inglés C2, Francés A2"),
        ],
        "training": [
            "Cursos de CompTIA Security+, CompTIA CySA+, COMP 601 Fundamentos de Ciberseguridad y Taller de MITRE ATT&CK y Respuesta a Incidentes.",
            "Práctica continua en seguridad defensiva, entrevistas de respuesta a incidentes, documentación técnica, análisis de comportamiento adversario y fundamentos de SOC.",
        ],
    },
}


def _escape(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def build_pdf(locale: Locale, content: ResumeContent, output: Path) -> None:
    from reportlab.lib.colors import black
    from reportlab.lib.enums import TA_CENTER, TA_RIGHT
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        BaseDocTemplate,
        Frame,
        HRFlowable,
        KeepTogether,
        ListFlowable,
        ListItem,
        PageTemplate,
        Paragraph,
        Spacer,
        Table,
        TableStyle,
    )

    name_style = ParagraphStyle(
        "name", fontName="Helvetica-Bold", fontSize=16.0, leading=17.2,
        alignment=TA_CENTER, textColor=black, spaceAfter=0.2,
    )
    title_style = ParagraphStyle(
        "title", fontName="Helvetica", fontSize=8.0, leading=9.2,
        alignment=TA_CENTER, textColor=black, spaceAfter=0.7,
    )
    contact_style = ParagraphStyle(
        "contact", fontName="Helvetica", fontSize=7.15, leading=8.3,
        alignment=TA_CENTER, textColor=black, spaceAfter=2.2,
    )
    section_style = ParagraphStyle(
        "section", fontName="Helvetica-Bold", fontSize=8.2, leading=9.2,
        textColor=black, spaceBefore=3.1, spaceAfter=0,
    )
    entry_left = ParagraphStyle(
        "entry-left", fontName="Helvetica", fontSize=8.0, leading=9.0,
        textColor=black, spaceAfter=0,
    )
    entry_right = ParagraphStyle(
        "entry-right", parent=entry_left, alignment=TA_RIGHT,
    )
    detail_style = ParagraphStyle(
        "detail", fontName="Helvetica-Oblique", fontSize=7.45, leading=8.35,
        textColor=black, leftIndent=0, spaceAfter=0.35,
    )
    bullet_style = ParagraphStyle(
        "bullet", fontName="Helvetica", fontSize=7.5, leading=8.6,
        textColor=black, spaceAfter=0.18,
    )
    skill_style = ParagraphStyle(
        "skill", fontName="Helvetica", fontSize=7.45, leading=8.55,
        textColor=black, spaceAfter=0.5,
    )

    def section_heading(label: str):
        return [
            Paragraph(_escape(label), section_style),
            HRFlowable(width="100%", thickness=0.55, color=black, spaceBefore=0.25, spaceAfter=1.4),
        ]

    def entry_block(entry: Entry):
        organization = f"<b>{_escape(entry['organization'])}</b>"
        location = _escape(entry["location"])
        first = Table(
            [[Paragraph(organization, entry_left), Paragraph(location, entry_right)]],
            colWidths=[5.55 * inch, 1.63 * inch],
        )
        role = f"<i>{_escape(entry['role'])}</i>"
        dates = f"<b>{_escape(entry['dates'])}</b>"
        second = Table(
            [[Paragraph(role, entry_left), Paragraph(dates, entry_right)]],
            colWidths=[5.55 * inch, 1.63 * inch],
        )
        table_style = TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 0),
            ("RIGHTPADDING", (0, 0), (-1, -1), 0),
            ("TOPPADDING", (0, 0), (-1, -1), 0),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
        ])
        first.setStyle(table_style)
        second.setStyle(table_style)

        block: list[object] = [first, second]
        if entry["detail"]:
            block.append(Paragraph(_escape(entry["detail"]), detail_style))
        if entry["bullets"]:
            items = [
                ListItem(Paragraph(_escape(bullet), bullet_style), leftIndent=7)
                for bullet in entry["bullets"]
            ]
            block.append(
                ListFlowable(
                    items,
                    bulletType="bullet",
                    start="circle",
                    leftIndent=11,
                    bulletFontName="Helvetica",
                    bulletFontSize=5.0,
                    bulletOffsetY=0.9,
                    spaceAfter=0.6,
                )
            )
        block.append(Spacer(1, 0.7))
        return KeepTogether(block)

    flow: list[object] = [
        Paragraph(NAME, name_style),
        Paragraph(_escape(content["title"]), title_style),
        Paragraph(_escape(content["contact"]), contact_style),
    ]

    for key in ("education", "work", "projects", "leadership"):
        flow.extend(section_heading(content["sections"][key]))
        for entry in content[key]:
            flow.append(entry_block(entry))

    flow.extend(section_heading(content["sections"]["skills"]))
    for label, value in content["skills"]:
        flow.append(Paragraph(f"<b>{_escape(label)}:</b> {_escape(value)}", skill_style))

    flow.extend(section_heading(content["sections"]["training"]))
    training_items = [
        ListItem(Paragraph(_escape(item), bullet_style), leftIndent=7)
        for item in content["training"]
    ]
    flow.append(
        ListFlowable(
            training_items,
            bulletType="bullet",
            start="circle",
            leftIndent=11,
            bulletFontName="Helvetica",
            bulletFontSize=5.0,
            bulletOffsetY=0.9,
        )
    )

    output.parent.mkdir(parents=True, exist_ok=True)
    doc = BaseDocTemplate(
        str(output),
        pagesize=letter,
        leftMargin=0.48 * inch,
        rightMargin=0.48 * inch,
        topMargin=0.36 * inch,
        bottomMargin=0.36 * inch,
        title=f"Jose Pablo Samano - {'Resume' if locale == 'en' else 'CV'}",
        author="Jose Pablo Samano Suarez",
    )
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="resume")
    doc.addPageTemplates([PageTemplate(id="resume", frames=[frame])])
    doc.build(flow)


def build_docx(locale: Locale, content: ResumeContent, output: Path) -> None:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    black = RGBColor(0, 0, 0)
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(8)
    normal.font.color.rgb = black
    normal.paragraph_format.space_after = Pt(0.5)
    normal.paragraph_format.line_spacing = 1.0

    for style_name in ("List Bullet", "List Paragraph"):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(7.7)
        style.font.color.rgb = black
        style.paragraph_format.left_indent = Inches(0.14)
        style.paragraph_format.first_line_indent = Inches(-0.10)
        style.paragraph_format.space_after = Pt(0.2)
        style.paragraph_format.line_spacing = 1.0

    section = doc.sections[0]
    section.start_type = WD_SECTION.CONTINUOUS
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.36)
    section.bottom_margin = Inches(0.36)
    section.left_margin = Inches(0.48)
    section.right_margin = Inches(0.48)
    section.header_distance = Inches(0.2)
    section.footer_distance = Inches(0.2)

    def set_font(run, size: float, bold: bool = False, italic: bool = False):
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = black

    def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_mar = tc_pr.first_child_found_in("w:tcMar")
        if tc_mar is None:
            tc_mar = OxmlElement("w:tcMar")
            tc_pr.append(tc_mar)
        for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
            node = tc_mar.find(qn(f"w:{name}"))
            if node is None:
                node = OxmlElement(f"w:{name}")
                tc_mar.append(node)
            node.set(qn("w:w"), str(value))
            node.set(qn("w:type"), "dxa")

    def remove_table_borders(table):
        tbl_pr = table._tbl.tblPr
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            element = OxmlElement(f"w:{edge}")
            element.set(qn("w:val"), "nil")
            borders.append(element)
        tbl_pr.append(borders)

    def add_section_heading(label: str):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(3)
        paragraph.paragraph_format.space_after = Pt(1.2)
        paragraph.paragraph_format.keep_with_next = True
        run = paragraph.add_run(label)
        set_font(run, 8.2, bold=True)
        p_pr = paragraph._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "000000")
        p_bdr.append(bottom)
        p_pr.append(p_bdr)

    def add_entry(entry: Entry):
        table = doc.add_table(rows=2, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        remove_table_borders(table)
        widths = (Inches(5.55), Inches(1.63))
        for row in table.rows:
            for index, cell in enumerate(row.cells):
                cell.width = widths[index]
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                set_cell_margins(cell)
                cell.paragraphs[0].paragraph_format.space_after = Pt(0)

        org = table.cell(0, 0).paragraphs[0]
        set_font(org.add_run(entry["organization"]), 8.0, bold=True)
        location = table.cell(0, 1).paragraphs[0]
        location.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_font(location.add_run(entry["location"]), 8.0)

        role = table.cell(1, 0).paragraphs[0]
        set_font(role.add_run(entry["role"]), 8.0, italic=True)
        dates = table.cell(1, 1).paragraphs[0]
        dates.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_font(dates.add_run(entry["dates"]), 8.0, bold=True)

        if entry["detail"]:
            detail = doc.add_paragraph()
            detail.paragraph_format.space_after = Pt(0.2)
            set_font(detail.add_run(entry["detail"]), 7.45, italic=True)

        for bullet in entry["bullets"]:
            paragraph = doc.add_paragraph(style="List Bullet")
            set_font(paragraph.add_run(bullet), 7.7)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(0)
    set_font(title.add_run(NAME), 16.0, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(0.5)
    set_font(subtitle.add_run(content["title"]), 8.0)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(1.5)
    set_font(contact.add_run(content["contact"]), 7.15)

    for key in ("education", "work", "projects", "leadership"):
        add_section_heading(content["sections"][key])
        for entry in content[key]:
            add_entry(entry)

    add_section_heading(content["sections"]["skills"])
    for label, value in content["skills"]:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0.3)
        set_font(paragraph.add_run(f"{label}: "), 7.7, bold=True)
        set_font(paragraph.add_run(value), 7.7)

    add_section_heading(content["sections"]["training"])
    for item in content["training"]:
        paragraph = doc.add_paragraph(style="List Bullet")
        set_font(paragraph.add_run(item), 7.7)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = f"Jose Pablo Samano - {'Resume' if locale == 'en' else 'CV'}"
    doc.core_properties.author = "Jose Pablo Samano Suarez"
    doc.save(output)


def build_resume(locale: Locale) -> None:
    content = CONTENT[locale]
    suffix = "en" if locale == "en" else "es"
    pdf_path = ROOT / "public" / "resume" / f"jp-samano-resume-{suffix}.pdf"
    docx_path = ROOT / "docs" / "resume" / f"jp-samano-resume-{suffix}.docx"
    output_pdf = ROOT / "output" / "pdf" / f"jp-samano-resume-{suffix}.pdf"

    build_pdf(locale, content, pdf_path)
    build_docx(locale, content, docx_path)
    output_pdf.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(pdf_path, output_pdf)

    if locale == "en":
        shutil.copy2(pdf_path, ROOT / "public" / "resume.pdf")

    print(pdf_path)
    print(docx_path)
    print(output_pdf)
