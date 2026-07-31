from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "resume" / "jp-samano-resume-en.docx"

BLUE = RGBColor(31, 77, 120)
INK = RGBColor(18, 32, 51)
MUTED = RGBColor(85, 95, 110)


def set_cell_text(cell, text, bold=False, color=INK):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(8.3)
    run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width):
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width.inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=50, start=90, bottom=50, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D8DEE8", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_rule(paragraph, color="C9D7EA"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.style = "ResumeHeading"
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    add_rule(p)
    return p


def add_role(doc, title, meta, bullets):
    p = doc.add_paragraph()
    p.style = "RoleTitle"
    p.paragraph_format.keep_with_next = True
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = INK
    if meta:
        m = p.add_run(f" | {meta}")
        m.font.color.rgb = MUTED
    for bullet in bullets:
        b = doc.add_paragraph(style="List Bullet")
        b.paragraph_format.left_indent = Inches(0.18)
        b.paragraph_format.first_line_indent = Inches(-0.18)
        b.paragraph_format.space_after = Pt(1.2)
        b.paragraph_format.line_spacing = 1.03
        run = b.add_run(bullet)
        run.font.size = Pt(8.6)
        run.font.color.rgb = INK


def set_document_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(8.3)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(1.8)
    normal.paragraph_format.line_spacing = 1.04

    styles = doc.styles

    heading = styles.add_style("ResumeHeading", 1)
    heading.font.name = "Calibri"
    heading.font.size = Pt(10.1)
    heading.font.color.rgb = BLUE
    heading.font.bold = True
    heading.paragraph_format.space_before = Pt(5.5)
    heading.paragraph_format.space_after = Pt(1.9)

    role = styles.add_style("RoleTitle", 1)
    role.font.name = "Calibri"
    role.font.size = Pt(8.9)
    role.paragraph_format.space_before = Pt(2)
    role.paragraph_format.space_after = Pt(1.2)

    for style_name in ["List Bullet", "List Paragraph"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(8.6)
        style.paragraph_format.left_indent = Inches(0.18)
        style.paragraph_format.first_line_indent = Inches(-0.18)
        style.paragraph_format.space_after = Pt(1.2)
        style.paragraph_format.line_spacing = 1.03


def build():
    doc = Document()
    set_document_styles(doc)

    section = doc.sections[0]
    section.start_type = WD_SECTION.CONTINUOUS
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.42)
    section.bottom_margin = Inches(0.42)
    section.left_margin = Inches(0.52)
    section.right_margin = Inches(0.52)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(0)
    r = title.add_run("JOSE PABLO SAMANO SUAREZ")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(18)
    r.font.color.rgb = INK

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(1.8)
    s = subtitle.add_run(
        "Computer Science Student | Software Engineering | Cybersecurity | Full-Stack Product Systems"
    )
    s.font.name = "Calibri"
    s.font.size = Pt(8.7)
    s.font.color.rgb = MUTED

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(4)
    c = contact.add_run(
        "Wolfville, NS / Mexico City | jpss2004@icloud.com | 6195978559 | "
        "linkedin.com/in/jose-pablo-samano-suarez | github.com/jpss2004-bot"
    )
    c.font.name = "Calibri"
    c.font.size = Pt(8)
    c.font.color.rgb = MUTED

    add_heading(doc, "PROFILE")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.add_run(
        "Computer Science student at Acadia University who ships production software for a paying client. Built and "
        "maintains two systems in daily use at LegalShelf: CheckWise, a REPSE compliance platform serving 3 client "
        "companies and 15+ vendors with close to 20,000 documents processed, and Verifaid, a document verification "
        "system. Works across full-stack engineering, domain modeling, and cybersecurity. Open to software engineering "
        "internships and junior technical roles."
    )

    add_heading(doc, "EDUCATION")
    add_role(
        doc,
        "Bachelor of Science, Computer Science, Acadia University",
        "Wolfville, NS | Expected May 2027",
        [
            "Focus: software engineering, data structures, databases, cybersecurity, systems analysis, human-machine interaction, and secure application design.",
        ],
    )
    add_role(
        doc,
        "High School Diploma, Trinity College School",
        "Port Hope, ON | June 2023",
        [],
    )

    add_heading(doc, "EXPERIENCE")
    add_role(
        doc,
        "Software Developer, LegalShelf",
        "Mexico City / Remote | 2026-Present",
        [
            "Builds and maintains two systems in production and daily use: CheckWise, a REPSE compliance platform, and Verifaid, a document verification system.",
            "CheckWise serves 3 client companies and 15+ vendors, with close to 20,000 compliance documents processed to date. Built the FastAPI backend and Next.js 15 / React 19 frontend covering evidence intake, reviewer adjudication, and portfolio-wide risk views with JWT auth and role-based access.",
            "Modeled REPSE obligations as institution x cycle (SAT, IMSS, INFONAVIT, STPS; monthly through annual) so the operating calendar and expediente gates derive from the data model, and shipped an AI-assisted reports centre guarded by an AI-safety suite inside 320+ backend tests.",
            "Built database cleanup and metadata-extraction pipelines that make large document sets consistent, queryable, and reliable to process downstream.",
        ],
    )
    add_role(
        doc,
        "Cybolt Academy Participant, Cybolt",
        "Mexico City | May-Sep 2025",
        [
            "Completed hands-on workshops on MITRE ATT&CK, ICS/OT security, cyber kill chain, incident response, and defensive blue-team operations.",
            "Analyzed adversary TTPs and ransomware behavior to connect attack patterns with detection, containment, and investigation, and built STAR-based incident-response interview cases.",
        ],
    )
    add_role(
        doc,
        "Production Operator, Michelin",
        "Waterville, NS | May-Sep 2024",
        [
            "Operated and monitored rubber-processing equipment in a high-pressure industrial environment, flagging delays and equipment issues to protect plant continuity.",
        ],
    )
    add_role(
        doc,
        "Data Processor, Legal Shelf",
        "Mexico City | May-Aug 2023",
        [
            "Digitized and organized legal documents with structured metadata under strict confidentiality and data-protection practices.",
        ],
    )

    add_heading(doc, "SELECTED PROJECTS")
    add_role(
        doc,
        "Founder / Developer, SAVR - Context-Aware Dining Recommendation Platform",
        "Wolfville, NS | 2025-Present",
        [
            "Deployed full-stack recommendation platform (FastAPI, React, TypeScript, SQL) that ranks venues using preferences, budget, social context, and user intent, and states why each result fits.",
            "Designed Describe Your Night, Build Your Night, and Surprise Me flows connecting structured inputs to prioritized, explainable recommendation cards. Live at context-aware-dining-platform-1.vercel.app.",
        ],
    )
    add_role(
        doc,
        "System Designer, ER Triage & Queue Manager",
        "2025",
        [
            "Modeled an emergency-room queue workflow around intake, vitals, ESI v4 acuity scoring, patient state, status history, and dashboard visibility using Python, NiceGUI, and SQLite, showing the clinical reasoning behind every assigned level.",
        ],
    )
    add_role(
        doc,
        "Developer, Family Phrase Game",
        "2026",
        [
            "Built and deployed a Flask web app that turns family-submitted phrases into a playable party game with phrase loading, scoring, and a simple live interface.",
        ],
    )

    add_heading(doc, "TECHNICAL SKILLS")
    table = doc.add_table(rows=5, cols=2)
    table.autofit = False
    set_table_borders(table)
    rows = [
        ("Software", "Python, Java, JavaScript/TypeScript, C, C#, SQL, React, Next.js, FastAPI, Flask, REST APIs, JWT auth, Git"),
        ("Data & Delivery", "PostgreSQL, SQLite, SQLAlchemy, Alembic, data modeling, metadata extraction, validation, testing, Vercel, Render"),
        ("Cybersecurity", "Incident response, threat analysis, MITRE ATT&CK, cyber kill chain, SOC fundamentals, ICS/OT security, authentication, access control"),
        ("Languages", "Spanish native, English C2, French A2"),
        ("Training", "CompTIA Security+ and CySA+ coursework, COMP 601 Cybersecurity Fundamentals, MITRE ATT&CK & Incident Response Workshop"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        set_cell_width(row.cells[0], Inches(1.28))
        set_cell_width(row.cells[1], Inches(6.18))
        set_cell_margins(row.cells[0])
        set_cell_margins(row.cells[1])
        set_cell_shading(row.cells[0], "E8EEF5")
        set_cell_text(row.cells[0], label, bold=True, color=BLUE)
        set_cell_text(row.cells[1], value)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
