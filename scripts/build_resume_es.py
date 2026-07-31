"""Build the Spanish resume (CV) in two formats from a single content source.

- PDF  -> public/resume/jp-samano-resume-es.pdf   (served by the site, ATS text-extractable)
- DOCX -> docs/resume/jp-samano-resume-es.docx    (editable Word source, mirrors the English doc)

Mirrors the structure of scripts/build_resume_docx.py (English). Run with the
project venv that has `reportlab` and `python-docx` installed:

    ~/.venvs/jp-resume/bin/python scripts/build_resume_es.py
"""

from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
PDF_OUT = ROOT / "public" / "resume" / "jp-samano-resume-es.pdf"
DOCX_OUT = ROOT / "docs" / "resume" / "jp-samano-resume-es.docx"

# Brand colors shared with the English resume.
BLUE = (31, 77, 120)
INK = (18, 32, 51)
MUTED = (85, 95, 110)

NAME = "JOSE PABLO SAMANO SUAREZ"
SUBTITLE = (
    "Estudiante de Ciencias de la Computación | Ingeniería de Software | "
    "Ciberseguridad | Sistemas de Producto Full-Stack"
)
CONTACT = (
    "Wolfville, NS / Ciudad de México | jpss2004@icloud.com | 6195978559 | "
    "linkedin.com/in/jose-pablo-samano-suarez | github.com/jpss2004-bot"
)

PROFILE = (
    "Estudiante de Ciencias de la Computación en Acadia University que entrega software "
    "en producción para un cliente real. Construyó y mantiene dos sistemas en uso diario "
    "en LegalShelf: CheckWise, plataforma de cumplimiento REPSE que atiende a 3 empresas "
    "cliente y 15+ proveedores con cerca de 20,000 documentos procesados, y Verifaid, un "
    "sistema de verificación de documentos. Trabaja en ingeniería full-stack, modelado de "
    "dominio y ciberseguridad. Abierto a prácticas y roles técnicos junior."
)

# Each section: (HEADING, [ (title, meta, [bullets]) , ... ])
ROLE_SECTIONS = [
    (
        "EDUCACIÓN",
        [
            (
                "Bachelor of Science, Computer Science, Acadia University",
                "Wolfville, NS | Esperado Mayo 2027",
                [
                    "Enfoque: ingeniería de software, estructuras de datos, bases de datos, ciberseguridad, análisis de sistemas, interacción humano-máquina y diseño de aplicaciones seguras.",
                ],
            ),
            (
                "High School Diploma, Trinity College School",
                "Port Hope, ON | Junio 2023",
                [],
            ),
        ],
    ),
    (
        "EXPERIENCIA",
        [
            (
                "Desarrollador de Software, LegalShelf",
                "Ciudad de México / Remoto | 2026-Presente",
                [
                    "Construye y mantiene dos sistemas en producción y uso diario: CheckWise, plataforma de cumplimiento REPSE, y Verifaid, sistema de verificación de documentos.",
                    "CheckWise atiende a 3 empresas cliente y 15+ proveedores, con cerca de 20,000 documentos de cumplimiento procesados a la fecha. Construyó el backend en FastAPI y el frontend en Next.js 15 / React 19 con recepción de evidencia, dictamen del revisor y vistas de riesgo del portafolio, con autenticación JWT y control de acceso por rol.",
                    "Modeló las obligaciones REPSE como institución x ciclo (SAT, IMSS, INFONAVIT, STPS; mensual a anual) para que el calendario operativo y las compuertas de expediente se deriven del modelo de datos, y entregó un centro de reportes con IA protegido por una suite de seguridad dentro de 320+ pruebas backend.",
                ],
            ),
            (
                "Participante de Cybolt Academy, Cybolt",
                "Ciudad de México | May-Sep 2025",
                [
                    "Completó talleres prácticos sobre MITRE ATT&CK, seguridad ICS/OT, cyber kill chain, respuesta a incidentes y operaciones defensivas de blue team.",
                    "Analizó TTPs adversarias y comportamiento de ransomware para conectar patrones de ataque con detección, contención e investigación, y construyó casos de entrevista basados en STAR.",
                ],
            ),
            (
                "Operador de Producción, Michelin",
                "Waterville, NS | May-Sep 2024",
                [
                    "Operó y monitoreó equipo de procesamiento de hule en un entorno industrial de alta presión, señalando retrasos y fallas de equipo para proteger la continuidad de la planta.",
                ],
            ),
            (
                "Procesador de Datos, Legal Shelf",
                "Ciudad de México | May-Ago 2023",
                [
                    "Digitalizó y organizó documentos legales con metadatos estructurados bajo estrictas prácticas de confidencialidad y protección de datos.",
                ],
            ),
        ],
    ),
    (
        "PROYECTOS SELECCIONADOS",
        [
            (
                "Desarrollador, Pipelines de Datos - Limpieza de Bases y Extracción de Metadatos",
                "2026",
                [
                    "Construyó pipelines de limpieza de bases de datos y extracción de metadatos para volver grandes conjuntos documentales consistentes, consultables y confiables de procesar.",
                ],
            ),
            (
                "Fundador / Desarrollador, SAVR - Plataforma de Recomendación Gastronómica Contextual",
                "Wolfville, NS | 2025-Presente",
                [
                    "Plataforma full-stack desplegada (FastAPI, React, TypeScript, SQL) que clasifica lugares según preferencias, presupuesto, contexto social e intención, y explica por qué encaja cada resultado.",
                    "Diseñó los flujos Describe Your Night, Build Your Night y Surprise Me. En vivo en context-aware-dining-platform-1.vercel.app.",
                ],
            ),
            (
                "Diseñador de Sistema, ER Triage & Queue Manager",
                "2025",
                [
                    "Modeló un flujo de cola de urgencias en torno a admisión, signos vitales, agudeza ESI v4, estado del paciente e historial usando Python, NiceGUI y SQLite, mostrando el razonamiento clínico de cada nivel asignado.",
                ],
            ),
            (
                "Desarrollador, Family Phrase Game",
                "2026",
                [
                    "Construyó y desplegó una aplicación web en Flask que convierte frases familiares en un juego de fiesta jugable, con carga de frases, puntuación e interfaz en vivo.",
                ],
            ),
        ],
    ),
]

SKILLS = [
    ("Software", "Python, Java, JavaScript/TypeScript, C, C#, SQL, React, Next.js, FastAPI, Flask, REST APIs, Git"),
    ("Datos y Entrega", "PostgreSQL, SQLite, SQLAlchemy, Alembic, modelado de datos, extracción de metadatos, validación, pruebas, Vercel, Render"),
    ("Ciberseguridad", "Respuesta a incidentes, análisis de amenazas, MITRE ATT&CK, cyber kill chain, fundamentos de SOC, seguridad ICS/OT, autenticación, control de acceso"),
    ("Idiomas", "Español nativo, Inglés C2, Francés A2"),
]

TRAINING = (
    "Seguridad y desarrollo profesional",
    "",
    [
        "Cursos de CompTIA Security+, cursos de CompTIA CySA+, COMP 601 Fundamentos de Ciberseguridad, Taller de MITRE ATT&CK y Respuesta a Incidentes.",
        "Práctica continua en seguridad defensiva, entrevistas de respuesta a incidentes, documentación técnica, análisis de comportamiento adversario y fundamentos de SOC.",
    ],
)


# --------------------------------------------------------------------------- PDF
def build_pdf():
    from reportlab.lib.colors import Color, HexColor
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        BaseDocTemplate,
        Frame,
        HRFlowable,
        ListFlowable,
        ListItem,
        PageTemplate,
        Paragraph,
        Spacer,
        Table,
        TableStyle,
    )
    from reportlab.lib.styles import ParagraphStyle

    def rl(rgb):
        return Color(rgb[0] / 255, rgb[1] / 255, rgb[2] / 255)

    blue, ink, muted = rl(BLUE), rl(INK), rl(MUTED)

    name_style = ParagraphStyle("name", fontName="Helvetica-Bold", fontSize=17,
                                leading=19, alignment=TA_CENTER, textColor=ink, spaceAfter=0)
    sub_style = ParagraphStyle("sub", fontName="Helvetica", fontSize=8.5, leading=10,
                               alignment=TA_CENTER, textColor=muted, spaceAfter=1.5)
    contact_style = ParagraphStyle("contact", fontName="Helvetica", fontSize=8, leading=9.6,
                                   alignment=TA_CENTER, textColor=muted, spaceAfter=2.5)
    heading_style = ParagraphStyle("heading", fontName="Helvetica-Bold", fontSize=9.8,
                                   leading=11, textColor=blue, spaceBefore=2.6, spaceAfter=0.6)
    body_style = ParagraphStyle("body", fontName="Helvetica", fontSize=8.1, leading=9.5,
                                textColor=ink, spaceAfter=2)
    role_style = ParagraphStyle("role", fontName="Helvetica", fontSize=8.3, leading=9.7,
                                textColor=ink, spaceBefore=1.6, spaceAfter=0.6)
    bullet_style = ParagraphStyle("bullet", fontName="Helvetica", fontSize=8.0, leading=9.2,
                                  textColor=ink, spaceAfter=0.8)
    skill_label = ParagraphStyle("sl", fontName="Helvetica-Bold", fontSize=8.0, leading=9.4,
                                 textColor=blue)
    skill_value = ParagraphStyle("sv", fontName="Helvetica", fontSize=8.0, leading=9.4,
                                 textColor=ink)

    def esc(text):
        return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    def rule():
        return HRFlowable(width="100%", thickness=0.8, color=HexColor("#C9D7EA"),
                          spaceBefore=0.5, spaceAfter=2)

    def role_block(title, meta, bullets):
        items = []
        head = f'<b>{esc(title)}</b>'
        if meta:
            head += f' <font color="#555F6E">| {esc(meta)}</font>'
        items.append(Paragraph(head, role_style))
        if bullets:
            items.append(ListFlowable(
                [ListItem(Paragraph(esc(b), bullet_style), leftIndent=10,
                          value="•", spaceb=1.6) for b in bullets],
                bulletType="bullet", bulletColor=ink, bulletFontSize=6,
                leftIndent=12, bulletOffsetY=0,
            ))
        return items

    flow = []
    flow.append(Paragraph(NAME, name_style))
    flow.append(Paragraph(esc(SUBTITLE), sub_style))
    flow.append(Paragraph(esc(CONTACT), contact_style))

    flow.append(Paragraph("PERFIL", heading_style))
    flow.append(rule())
    flow.append(Paragraph(esc(PROFILE), body_style))

    for heading, roles in ROLE_SECTIONS:
        flow.append(Paragraph(heading, heading_style))
        flow.append(rule())
        for title, meta, bullets in roles:
            flow.extend(role_block(title, meta, bullets))

    flow.append(Paragraph("HABILIDADES TÉCNICAS", heading_style))
    flow.append(rule())
    data = [[Paragraph(esc(lbl), skill_label), Paragraph(esc(val), skill_value)]
            for lbl, val in SKILLS]
    table = Table(data, colWidths=[1.28 * inch, 6.12 * inch])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, -1), HexColor("#E8EEF5")),
        ("GRID", (0, 0), (-1, -1), 0.4, HexColor("#D8DEE8")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 2.5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 2.5),
    ]))
    flow.append(table)

    flow.append(Paragraph("FORMACIÓN ADICIONAL", heading_style))
    flow.append(rule())
    flow.extend(role_block(*TRAINING))

    PDF_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = BaseDocTemplate(
        str(PDF_OUT), pagesize=letter,
        leftMargin=0.52 * inch, rightMargin=0.52 * inch,
        topMargin=0.42 * inch, bottomMargin=0.42 * inch,
        title="Jose Pablo Samano - CV", author="Jose Pablo Samano Suarez",
    )
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="main")
    doc.addPageTemplates([PageTemplate(id="main", frames=[frame])])
    doc.build(flow)
    print(PDF_OUT)


# -------------------------------------------------------------------------- DOCX
def build_docx():
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    blue = RGBColor(*BLUE)
    ink = RGBColor(*INK)
    muted = RGBColor(*MUTED)

    def add_rule(paragraph, color="C9D7EA"):
        p_pr = paragraph._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:space"), "3")
        bottom.set(qn("w:color"), color)
        p_bdr.append(bottom)
        p_pr.append(p_bdr)

    def add_heading(doc, text):
        p = doc.add_paragraph()
        p.style = "ResumeHeading"
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        add_rule(p)

    def add_role(doc, title, meta, bullets):
        p = doc.add_paragraph()
        p.style = "RoleTitle"
        p.paragraph_format.keep_with_next = True
        r = p.add_run(title)
        r.bold = True
        r.font.color.rgb = ink
        if meta:
            m = p.add_run(f" | {meta}")
            m.font.color.rgb = muted
        for bullet in bullets:
            b = doc.add_paragraph(style="List Bullet")
            b.paragraph_format.left_indent = Inches(0.18)
            b.paragraph_format.first_line_indent = Inches(-0.18)
            b.paragraph_format.space_after = Pt(1.6)
            b.paragraph_format.line_spacing = 1.03
            run = b.add_run(bullet)
            run.font.size = Pt(8.6)
            run.font.color.rgb = ink

    def set_cell_text(cell, text, bold=False, color=ink):
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        run.bold = bold
        run.font.name = "Calibri"
        run.font.size = Pt(8.7)
        run.font.color.rgb = color

    def set_cell_shading(cell, fill):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)

    def set_table_borders(table, color="D8DEE8", size="4"):
        tbl_pr = table._tbl.tblPr
        borders = OxmlElement("w:tblBorders")
        for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), size)
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color)
            borders.append(el)
        tbl_pr.append(borders)

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(8.7)
    normal.font.color.rgb = ink
    normal.paragraph_format.space_after = Pt(2.2)
    normal.paragraph_format.line_spacing = 1.04

    heading = doc.styles.add_style("ResumeHeading", 1)
    heading.font.name = "Calibri"
    heading.font.size = Pt(10.1)
    heading.font.color.rgb = blue
    heading.font.bold = True
    heading.paragraph_format.space_before = Pt(5.5)
    heading.paragraph_format.space_after = Pt(2.3)

    role = doc.styles.add_style("RoleTitle", 1)
    role.font.name = "Calibri"
    role.font.size = Pt(8.9)
    role.paragraph_format.space_before = Pt(2)
    role.paragraph_format.space_after = Pt(1.2)

    section = doc.sections[0]
    section.start_type = WD_SECTION.CONTINUOUS
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.42)
    section.bottom_margin = Inches(0.42)
    section.left_margin = Inches(0.52)
    section.right_margin = Inches(0.52)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(0)
    r = title.add_run(NAME)
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(18)
    r.font.color.rgb = ink

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(1.8)
    s = subtitle.add_run(SUBTITLE)
    s.font.name = "Calibri"
    s.font.size = Pt(8.7)
    s.font.color.rgb = muted

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(4)
    c = contact.add_run(CONTACT)
    c.font.name = "Calibri"
    c.font.size = Pt(8)
    c.font.color.rgb = muted

    add_heading(doc, "PERFIL")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.add_run(PROFILE)

    for hdg, roles in ROLE_SECTIONS:
        add_heading(doc, hdg)
        for t, meta, bullets in roles:
            add_role(doc, t, meta, bullets)

    add_heading(doc, "HABILIDADES TÉCNICAS")
    table = doc.add_table(rows=len(SKILLS), cols=2)
    table.autofit = False
    set_table_borders(table)
    for row, (label, value) in zip(table.rows, SKILLS):
        row.cells[0].width = Inches(1.28)
        row.cells[1].width = Inches(6.18)
        set_cell_shading(row.cells[0], "E8EEF5")
        set_cell_text(row.cells[0], label, bold=True, color=blue)
        set_cell_text(row.cells[1], value)

    add_heading(doc, "FORMACIÓN ADICIONAL")
    add_role(doc, *TRAINING)

    DOCX_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_OUT)
    print(DOCX_OUT)


if __name__ == "__main__":
    build_pdf()
    build_docx()
